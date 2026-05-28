"""Main document processor for handling multiple file types."""

import os
from pathlib import Path
from typing import List, Dict, Any
import logging

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

from app.utils.logger import get_logger
from app.utils.validators import validate_file, validate_document_size
from .text_chunker import TextChunker


logger = get_logger(__name__)


class DocumentProcessor:
    """Process and extract text from various document types."""

    SUPPORTED_FORMATS = {".pdf": "pdf", ".txt": "text", ".docx": "docx", ".doc": "docx"}

    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        max_upload_size_mb: int = 100,
    ):
        """
        Initialize document processor.

        Args:
            chunk_size: Size of text chunks
            chunk_overlap: Overlap between chunks
            max_upload_size_mb: Maximum upload size in MB
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.max_upload_size_mb = max_upload_size_mb
        self.chunker = TextChunker(chunk_size, chunk_overlap)

    def process_document(self, file_path: str) -> Dict[str, Any]:
        """
        Process a document and extract text with metadata.

        Args:
            file_path: Path to document file

        Returns:
            Dictionary with extracted text and metadata
        """
        # Validate file
        if not validate_file(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        if not validate_document_size(file_path, self.max_upload_size_mb):
            raise ValueError(
                f"File size exceeds maximum allowed size of {self.max_upload_size_mb}MB"
            )

        file_ext = Path(file_path).suffix.lower()
        if file_ext not in self.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported file format: {file_ext}")

        # Extract text based on file type
        if file_ext == ".pdf":
            text = self._extract_pdf_text(file_path)
        elif file_ext == ".docx":
            text = self._extract_docx_text(file_path)
        elif file_ext == ".txt":
            text = self._extract_text_file(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")

        # Create chunks
        chunks = self.chunker.chunk_text(text)

        # Prepare metadata
        metadata = {
            "filename": Path(file_path).name,
            "file_type": file_ext[1:],
            "total_characters": len(text),
            "total_chunks": len(chunks),
            "chunk_size": self.chunk_size,
            "processed_at": str(Path(file_path).stat().st_mtime),
        }

        logger.info(f"Processed document: {metadata['filename']} - {len(chunks)} chunks")

        return {
            "text": text,
            "chunks": chunks,
            "metadata": metadata,
        }

    def process_multiple_documents(self, directory_path: str) -> List[Dict[str, Any]]:
        """
        Process all documents in a directory.

        Args:
            directory_path: Path to directory containing documents

        Returns:
            List of processed documents
        """
        results = []
        directory = Path(directory_path)

        if not directory.exists():
            raise FileNotFoundError(f"Directory not found: {directory_path}")

        for file_path in directory.iterdir():
            if file_path.suffix.lower() in self.SUPPORTED_FORMATS:
                try:
                    result = self.process_document(str(file_path))
                    results.append(result)
                except Exception as e:
                    logger.error(f"Error processing {file_path}: {str(e)}")
                    continue

        logger.info(f"Processed {len(results)} documents from {directory_path}")
        return results

    @staticmethod
    def _extract_pdf_text(file_path: str) -> str:
        """Extract text from PDF file."""
        if not PyPDF2:
            raise ImportError("PyPDF2 is required to process PDF files")

        text = []
        try:
            with open(file_path, "rb") as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text.append(page.extract_text())
        except Exception as e:
            logger.error(f"Error extracting PDF text: {str(e)}")
            raise

        return "\n".join(text)

    @staticmethod
    def _extract_docx_text(file_path: str) -> str:
        """Extract text from DOCX file."""
        if not DocxDocument:
            raise ImportError("python-docx is required to process DOCX files")

        text = []
        try:
            doc = DocxDocument(file_path)
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text.append(cell.text)
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {str(e)}")
            raise

        return "\n".join(text)

    @staticmethod
    def _extract_text_file(file_path: str) -> str:
        """Extract text from TXT file."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error extracting text file: {str(e)}")
            raise
